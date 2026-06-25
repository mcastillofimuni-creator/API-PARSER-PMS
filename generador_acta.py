import os
import re
import tempfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document


TEMPLATE_NAME = "PLANTILLA_ACTA_INTERFERENCIAS_BASE.docx"


def _root_dir() -> Path:
    return Path(__file__).resolve().parent


def _template_path() -> Path:
    path = _root_dir() / TEMPLATE_NAME
    if not path.exists():
        raise FileNotFoundError(
            f"No se encontró la plantilla {TEMPLATE_NAME}. "
            "Debe estar en la misma carpeta que main.py."
        )
    return path


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _normalizar_espacios(texto: str) -> str:
    return re.sub(r"\s+", " ", _safe_text(texto)).strip()




ACRONIMOS = {
    "PMS", "OT", "SAP", "HSE&Q", "HSEQ", "HSE", "I&C", "CT", "C.T.",
    "JM", "JMI", "T&D", "KVSG", "SGS", "S.A.", "SAC", "S.A.C.", "S.A", "SRL", "S.R.L."
}

EMPRESA_MAP = {
    "MAGNEX": "Magnex",
    "JM INGENIEROS": "JM Ingenieros",
    "JMI": "JMI",
    "JIMSAO": "Jimsao",
    "GERER L ENERGIE SAC": "Gerer L Energie S.A.C.",
    "ULLOA": "Ulloa",
    "ULLOA S.A": "Ulloa S.A.",
    "ULLOA S.A.": "Ulloa S.A.",
    "SEFREL": "Sefrel",
    "UNITELEC": "Unitelec",
    "DIM": "DIM",
    "T&D ELECTRIC": "T&D Electric",
    "MAQUIRENTAS": "Maquirentas",
    "MAQUIRENTA": "Maquirenta",
    "ORYGEN": "Orygen",
}

PALABRAS_MAP = {
    "hseq": "HSE&Q",
    "hse&q": "HSE&Q",
    "hsq": "HSE&Q",
    "hse": "HSE",
    "orygen": "Orygen",
    "mantto": "mantenimiento",
    "pms": "PMS",
    "sap": "SAP",
    "ot": "OT",
}


def _formato_nombre(texto: Any) -> str:
    """Nombres propios en modo oración: Manuel Castillo, no MANUEL CASTILLO."""
    t = _normalizar_espacios(texto)
    if not t:
        return ""
    partes = []
    for palabra in t.split(" "):
        if not palabra:
            continue
        up = palabra.upper()
        if up in {"DE", "DEL", "LA", "LAS", "LOS", "Y"}:
            partes.append(up.lower())
        else:
            partes.append(palabra[:1].upper() + palabra[1:].lower())
    return " ".join(partes)


def _formato_empresa(texto: Any) -> str:
    """Empresa legible y consistente, conservando acrónimos razonables."""
    t = _normalizar_espacios(texto)
    if not t:
        return ""
    key = t.upper().replace(".", ".")
    key_simple = re.sub(r"\s+", " ", key).strip()
    if key_simple in EMPRESA_MAP:
        return EMPRESA_MAP[key_simple]

    palabras = []
    for palabra in t.split(" "):
        limpia = palabra.strip()
        up = limpia.upper().strip()
        if up in EMPRESA_MAP:
            palabras.append(EMPRESA_MAP[up])
        elif up in ACRONIMOS or len(up) <= 3 and up.isalpha():
            palabras.append(up)
        elif up == "SAC":
            palabras.append("S.A.C.")
        elif up == "SA" or up == "S.A":
            palabras.append("S.A.")
        else:
            palabras.append(limpia[:1].upper() + limpia[1:].lower())
    return " ".join(palabras)


def _corregir_texto_basico(texto: Any) -> str:
    """Corrección simple sin IA: espacios, abreviaturas típicas y mayúsculas excesivas."""
    t = _safe_text(texto)
    if not t:
        return ""
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    t = re.sub(r"[ \t]+", " ", t)

    # Correcciones comunes.
    reemplazos = {
        r"\bhseq\b": "HSE&Q",
        r"\bhse&q\b": "HSE&Q",
        r"\bhsq\b": "HSE&Q",
        r"\bmantto\b": "mantenimiento",
        r"\borygen\b": "Orygen",
        r"\bpms\b": "PMS",
        r"\bsap\b": "SAP",
    }
    for pat, rep in reemplazos.items():
        t = re.sub(pat, rep, t, flags=re.IGNORECASE)

    # Si una línea completa viene en mayúsculas, pasarla a formato nombre/empresa aproximado.
    lineas = []
    for raw in t.split("\n"):
        line = raw.strip()
        if not line:
            lineas.append("")
            continue
        if len(line) > 3 and line.upper() == line and re.search(r"[A-ZÁÉÍÓÚÑ]", line):
            # No tocar acrónimos cortos puros.
            line = " ".join(
                w if w in {"PMS", "SAP", "OT", "HSE&Q", "HSE"} else (w[:1].upper() + w[1:].lower())
                for w in line.split(" ")
            )
        lineas.append(line)
    return "\n".join(lineas).strip()


def _replace_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    """
    Reemplazo conservador de placeholders dentro de párrafos.
    Si el placeholder está partido en varios runs, reconstruye el párrafo.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace(key, value)

    if new_text == full_text:
        return

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def _replace_in_cell(cell, replacements: Dict[str, str]) -> None:
    for p in cell.paragraphs:
        _replace_in_paragraph(p, replacements)

    for table in cell.tables:
        for row in table.rows:
            for c in row.cells:
                _replace_in_cell(c, replacements)


def _replace_all(doc: Document, replacements: Dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, replacements)

    for section in doc.sections:
        for container in [section.header, section.footer]:
            for p in container.paragraphs:
                _replace_in_paragraph(p, replacements)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _replace_in_cell(cell, replacements)


def _set_cell_text(cell, text: str) -> None:
    """Escribe texto conservando, en lo posible, el estilo del primer run de la celda."""
    value = _safe_text(text)
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = value
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.add_run(value)
        for extra_p in cell.paragraphs[1:]:
            for run in extra_p.runs:
                run.text = ""
    else:
        cell.text = value


def _copiar_estilo_fila(row_origen, row_destino) -> None:
    """
    Copia propiedades de fila y celdas para que las filas agregadas mantengan
    el formato visual de la plantilla.
    """
    row_destino._tr.get_or_add_trPr().clear()
    for child in row_origen._tr.get_or_add_trPr():
        row_destino._tr.get_or_add_trPr().append(deepcopy(child))

    for src_cell, dst_cell in zip(row_origen.cells, row_destino.cells):
        src_tcpr = src_cell._tc.get_or_add_tcPr()
        dst_tcpr = dst_cell._tc.get_or_add_tcPr()
        dst_tcpr.clear()
        for child in src_tcpr:
            dst_tcpr.append(deepcopy(child))


def _asegurar_filas(tabla, cantidad_total: int, fila_modelo_idx: int = 1) -> None:
    while len(tabla.rows) < cantidad_total:
        new_row = tabla.add_row()
        try:
            _copiar_estilo_fila(tabla.rows[fila_modelo_idx], new_row)
        except Exception:
            pass


def _limpiar_filas_participantes(tabla) -> None:
    # Mantener encabezado. Limpiar filas de datos.
    for idx in range(1, len(tabla.rows)):
        row = tabla.rows[idx]
        for cidx, cell in enumerate(row.cells):
            _set_cell_text(cell, str(idx) if cidx == 0 else "")


def _normalizar_separadores_participante(texto: str) -> str:
    """Unifica separadores raros que suelen venir desde Word/Copilot."""
    t = _safe_text(texto)
    # Guiones: hyphen, en dash, em dash, non-breaking hyphen, minus sign.
    t = re.sub(r"[\u2010\u2011\u2012\u2013\u2014\u2212]", "-", t)
    t = t.replace("：", ":")
    return t


def _normalizar_area_participante(area: str, contexto_orygen: bool = False) -> str:
    a = _normalizar_espacios(area)
    if not a:
        return "Orygen" if contexto_orygen else ""

    au = a.upper()
    if "HSE" in au or "HSQ" in au or "HSEQ" in au:
        return "HSE&Q"
    if "MANT" in au:
        return "Mantenimiento"
    if "OPER" in au:
        return "Operaciones"
    if "ORYGEN" in au:
        return "Orygen"
    return _formato_empresa(a)



def _es_area_participante(texto: str) -> bool:
    """Reconoce si un texto parece ser un área/rol y no un nombre."""
    t = _normalizar_espacios(texto).upper()
    if not t:
        return False
    t = t.replace("Á", "A").replace("É", "E").replace("Í", "I").replace("Ó", "O").replace("Ú", "U")
    claves = [
        "HSE", "HSEQ", "HSE&Q", "HSQ",
        "MANT", "MANTTO", "MANTENIMIENTO", "MANTENIMENTO",
        "OPER", "OPERACIONES", "ORYGEN",
        "ELECT", "ELECTRICIDAD", "MECAN", "MECANICA",
        "INSTRUMENT", "I&C", "CONTROL",
    ]
    return any(k in t for k in claves)

def _extraer_participantes_adicionales(notas: str) -> List[Dict[str, str]]:
    """
    Parser simple y robusto de notas. Detecta formatos como:
    - Hector Tinoco (Orygen - HSE&Q)
    - Hector Tinoco (HSEq)
    - Hector Tinoco – HSE&Q
    - Hector Tinoco — HSE&Q
    - Hector Tinoco - HSE&Q
    - Hector Tinoco: HSE&Q
    - Mantenimiento: Manuel Castillo / Miguel Tasayco
    - Mantto: Manuel Castillo y Miguel Tasayco
    - HSE&Q: Hector Tinoco, Carolina Mostajo

    También interpreta el bloque posterior a frases tipo:
    "Participaron por parte de Orygen:".
    """
    participantes = []
    vistos = set()
    en_bloque_orygen = False

    def agregar(nombre: str, empresa: str):
        nombre = _normalizar_espacios(nombre)
        empresa = _normalizar_area_participante(empresa, contexto_orygen=en_bloque_orygen)
        if not nombre:
            return
        # Evitar frases que no son nombres.
        if nombre.upper().startswith(("PARTICIPARON", "PARTICIPANTES", "COMENTARIOS", "NOTA", "OBSERVACIONES", "ACUERDOS")):
            return
        # Evitar que un área sea interpretada como nombre.
        if _es_area_participante(nombre):
            return
        key = (nombre.upper(), empresa.upper())
        if key in vistos:
            return
        participantes.append({"nombre": _formato_nombre(nombre), "contrato": "", "empresa": empresa or "Orygen"})
        vistos.add(key)

    def dividir_nombres(texto_nombres: str) -> List[str]:
        texto_nombres = _normalizar_espacios(texto_nombres)
        if not texto_nombres:
            return []
        # Separadores típicos para varios participantes.
        partes = re.split(r"\s*(?:/|,|;|\by\b|\+)\s*", texto_nombres, flags=re.IGNORECASE)
        return [_normalizar_espacios(x) for x in partes if _normalizar_espacios(x)]

    for raw in _corregir_texto_basico(notas).splitlines():
        line = _normalizar_separadores_participante(raw).strip(" -*\t")
        if not line:
            continue

        line_upper = line.upper()
        if "PARTICIP" in line_upper and ("ORYGEN" in line_upper or "PARTE" in line_upper):
            en_bloque_orygen = True
            continue

        # Formato invertido: Área: Nombre / Nombre
        # Ej.: Mantenimiento: Manuel Castillo/Miguel Tasayco
        # Ej.: HSE&Q: Hector Tinoco, Carolina Mostajo
        m = re.match(
            r"^([A-Za-zÁÉÍÓÚÜÑáéíóúüñ &./]+?)\s*:\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ .'/,+;-]{3,})\s*$",
            line,
        )
        if m and _es_area_participante(m.group(1)):
            area = m.group(1)
            for nombre in dividir_nombres(m.group(2)):
                agregar(nombre, area)
            continue

        # Formato invertido con guion: Área - Nombre / Nombre
        # Ej.: Mantenimiento - Manuel Castillo / Miguel Tasayco
        m = re.match(
            r"^([A-Za-zÁÉÍÓÚÜÑáéíóúüñ &./]+?)\s*-\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ .'/,+;-]{3,})\s*$",
            line,
        )
        if m and _es_area_participante(m.group(1)):
            area = m.group(1)
            for nombre in dividir_nombres(m.group(2)):
                agregar(nombre, area)
            continue

        # Nombre (Orygen - Área) o Nombre (Área)
        m = re.match(r"^([A-Za-zÁÉÍÓÚÜÑáéíóúüñ .']{3,})\s*\(([^)]+)\)\s*$", line)
        if m:
            nombre = m.group(1)
            info = _normalizar_separadores_participante(m.group(2))
            partes = [x.strip() for x in re.split(r"\s*-\s*", info) if x.strip()]
            area = partes[-1] if partes else info
            agregar(nombre, area)
            continue

        # Nombre - Área / Nombre: Área.
        m = re.match(
            r"^([A-Za-zÁÉÍÓÚÜÑáéíóúüñ .']{3,}?)\s*(?:-|:|;)\s*([A-Za-z0-9ÁÉÍÓÚÜÑáéíóúüñ &./]+)\s*$",
            line,
        )
        if m:
            agregar(m.group(1), m.group(2))
            continue

        # En bloque Orygen, aceptar una línea con solo nombre como participante Orygen.
        if en_bloque_orygen and re.match(r"^[A-Za-zÁÉÍÓÚÜÑáéíóúüñ .']{5,}$", line):
            agregar(line, "Orygen")

    return participantes

def _estado_simple(estado: str) -> str:
    e = _safe_text(estado).upper()
    if not e:
        return "pendiente de validación"
    if "CONFORME" in e and "OBS" not in e:
        return "conforme"
    if "SIN ARCHIVO" in e:
        return "sin archivo"
    if "OBSERVADO" in e or "CORRECCIÓN" in e or "CORRECCION" in e:
        return "con observaciones"
    if "ERROR" in e:
        return "con error de validación"
    return e.lower()


def _armar_observaciones(payload: Dict[str, Any]) -> str:
    notas = _corregir_texto_basico(payload.get("notas"))
    empresas = payload.get("empresas") or []
    faltantes = payload.get("faltantes") or []

    total = len(empresas)
    observadas = [
        e for e in empresas
        if any(k in _safe_text(e.get("estado_validacion")).upper() for k in ["OBSERVADO", "CORRECCIÓN", "CORRECCION", "ERROR", "SIN ARCHIVO"])
    ]

    lineas = []
    lineas.append(
        f"Se efectuó la reunión de coordinación e identificación de interferencias correspondiente al {payload.get('pms') or 'PMS'}, "
        f"para la semana {payload.get('rango_semana') or payload.get('semana') or ''}, en la instalación {payload.get('central') or ''}."
    )

    if total:
        lineas.append(
            f"Se revisó la programación semanal registrada en la plataforma, con participación de {total} empresa(s) proveedora(s)."
        )

    if observadas:
        resumen_obs = ", ".join(
            f"{_safe_text(e.get('empresa'))} ({_estado_simple(e.get('estado_validacion'))})"
            for e in observadas[:10]
        )
        lineas.append(f"Se identificaron programas que requieren seguimiento o revisión: {resumen_obs}.")
    else:
        lineas.append("No se identificaron observaciones relevantes en la documentación revisada.")

    if faltantes:
        lineas.append(f"Empresas pendientes de registro o carga de programa: {', '.join(map(str, faltantes))}.")

    if notas:
        lineas.append("")
        lineas.append("Comentarios relevantes registrados:")
        lineas.append(notas)

    # Mantener una frase estándar útil para el formato.
    if "interferencia" not in notas.lower():
        lineas.append("")
        lineas.append("No se identificaron interferencias relevantes adicionales durante la reunión, salvo las que pudieran derivarse de la coordinación operativa de las actividades programadas.")

    return "\n".join([l for l in lineas if l is not None]).strip()


def _armar_acciones(payload: Dict[str, Any]) -> List[Dict[str, str]]:
    acciones = list(payload.get("acciones") or [])
    empresas = payload.get("empresas") or []

    if acciones:
        return [
            {
                "accion": _safe_text(a.get("accion") or a.get("descripcion")),
                "responsable": _safe_text(a.get("responsable")),
            }
            for a in acciones
            if _safe_text(a.get("accion") or a.get("descripcion"))
        ]

    generadas = []

    observadas = [
        e for e in empresas
        if any(k in _safe_text(e.get("estado_validacion")).upper() for k in ["OBSERVADO", "CORRECCIÓN", "CORRECCION", "ERROR"])
    ]
    sin_archivo = [
        e for e in empresas
        if "SIN ARCHIVO" in _safe_text(e.get("estado_validacion")).upper()
    ]

    for e in observadas[:5]:
        empresa = _formato_empresa(e.get("empresa"))
        generadas.append({
            "accion": f"Regularizar las observaciones identificadas en el programa PMS presentado por {empresa}.",
            "responsable": empresa,
        })

    for e in sin_archivo[:3]:
        empresa = _formato_empresa(e.get("empresa"))
        generadas.append({
            "accion": f"Completar la carga del archivo de programa semanal correspondiente a {empresa}.",
            "responsable": empresa,
        })

    faltantes = payload.get("faltantes") or []
    for empresa in faltantes[:5]:
        empresa = _formato_empresa(empresa)
        if empresa and not any(a.get("responsable", "").upper() == empresa.upper() for a in generadas):
            generadas.append({
                "accion": f"Coordinar con {empresa} la carga o regularización de su programa semanal.",
                "responsable": empresa,
            })

    if not generadas:
        generadas.append({
            "accion": "Mantener la coordinación semanal de actividades y comunicar oportunamente cualquier interferencia, restricción o cambio en la programación.",
            "responsable": "PROVEEDORES / ORYGEN",
        })

    return generadas


def generar_acta_interferencias(payload: Dict[str, Any]) -> Dict[str, str]:
    """
    Genera un DOCX usando la plantilla oficial RG02-P.HS.PE.013.
    No usa IA: rellena el formato con reglas y datos estructurados.
    """
    doc = Document(str(_template_path()))

    fecha = _safe_text(payload.get("fecha_reunion"))
    if not fecha:
        fecha = datetime.now().strftime("%d.%m.%Y")

    central = _safe_text(payload.get("central") or payload.get("central_label"))
    if normalizar := central.upper().replace(" ", ""):
        if "SANTAROSA" in normalizar:
            central = "C. T. Santa Rosa"
        elif "VENTANILLA" in normalizar:
            central = "C. C. Ventanilla"
        else:
            central = _formato_empresa(central)
    observaciones = _armar_observaciones(payload)
    acciones = [
        {
            "accion": _corregir_texto_basico(a.get("accion", "")),
            "responsable": _formato_empresa(a.get("responsable", "")),
        }
        for a in _armar_acciones(payload)
    ]

    replacements = {
        "[[FECHA_REUNION]]": fecha,
        "[[CENTRAL]]": central,
        "[[OBSERVACIONES]]": observaciones,
        "[[ACCION_01]]": acciones[0]["accion"] if len(acciones) > 0 else "",
        "[[RESPONSABLE_01]]": acciones[0]["responsable"] if len(acciones) > 0 else "",
        "[[ACCION_02]]": acciones[1]["accion"] if len(acciones) > 1 else "",
        "[[RESPONSABLE_02]]": acciones[1]["responsable"] if len(acciones) > 1 else "",
        # Compatibilidad con plantillas antiguas que todavía traen valores fijos.
        "17.06.2026": fecha,
        "17/06/2026": fecha,
        "CT SANTA ROSA": central,
        "CT VENTANILLA": central,
    }
    _replace_all(doc, replacements)

    # Participantes: tabla 1
    participantes = []
    vistos = set()

    for e in payload.get("empresas") or []:
        nombre = _safe_text(e.get("expositor"))
        empresa = _safe_text(e.get("empresa"))
        contrato = _safe_text(e.get("contrato"))
        if not nombre and not empresa:
            continue
        key = (nombre.upper(), empresa.upper())
        if key in vistos:
            continue
        participantes.append({"nombre": _formato_nombre(nombre), "contrato": contrato, "empresa": _formato_empresa(empresa)})
        vistos.add(key)

    for p in payload.get("participantes_adicionales") or []:
        nombre = _safe_text(p.get("nombre"))
        empresa = _safe_text(p.get("empresa") or "ORYGEN")
        contrato = _safe_text(p.get("contrato"))
        if not nombre:
            continue
        key = (nombre.upper(), empresa.upper())
        if key in vistos:
            continue
        participantes.append({"nombre": _formato_nombre(nombre), "contrato": contrato, "empresa": _formato_empresa(empresa)})
        vistos.add(key)

    for p in _extraer_participantes_adicionales(payload.get("notas") or ""):
        key = (p["nombre"].upper(), p["empresa"].upper())
        if key in vistos:
            continue
        participantes.append(p)
        vistos.add(key)

    if len(doc.tables) >= 2:
        tabla_part = doc.tables[1]
        _asegurar_filas(tabla_part, max(15, len(participantes) + 1), fila_modelo_idx=1)
        _limpiar_filas_participantes(tabla_part)

        for idx, p in enumerate(participantes, start=1):
            row = tabla_part.rows[idx]
            _set_cell_text(row.cells[0], str(idx))
            _set_cell_text(row.cells[1], _formato_nombre(p.get("nombre", "")))
            _set_cell_text(row.cells[2], p.get("contrato", ""))
            _set_cell_text(row.cells[3], _formato_empresa(p.get("empresa", "")))
            _set_cell_text(row.cells[4], "")

    # Acciones extra: tabla 2, filas 3 en adelante.
    if len(doc.tables) >= 3:
        tabla_acc = doc.tables[2]
        primera_fila_accion = 3
        filas_necesarias = primera_fila_accion + max(len(acciones), 2)
        _asegurar_filas(tabla_acc, filas_necesarias, fila_modelo_idx=3)

        for i in range(primera_fila_accion, len(tabla_acc.rows)):
            for c in tabla_acc.rows[i].cells:
                _set_cell_text(c, "")

        for i, accion in enumerate(acciones):
            row = tabla_acc.rows[primera_fila_accion + i]
            _set_cell_text(row.cells[0], accion.get("accion", ""))
            _set_cell_text(row.cells[1], accion.get("responsable", ""))

    out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    out.close()

    doc.save(out.name)

    pms = _safe_text(payload.get("pms")).replace(" ", "_") or "PMS"
    central_file = re.sub(r"[^A-Z0-9]+", "_", central.upper()).strip("_") or "CENTRAL"
    fecha_file = fecha.replace("/", ".").replace("-", ".")
    nombre_archivo = f"ACTA_INTERFERENCIAS_{pms}_{central_file}_{fecha_file}.docx"

    return {
        "archivo_generado": out.name,
        "nombre_archivo": nombre_archivo,
    }
